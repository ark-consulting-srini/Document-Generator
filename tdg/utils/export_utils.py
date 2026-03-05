"""
Export Utilities for T - BRD STTM Generator

Functions for exporting data to Excel, Word, and other formats.
"""

import io
import pandas as pd
from typing import Dict, List, Optional


def create_excel_export(dataframes: Dict[str, pd.DataFrame], 
                       include_brd: bool = False,
                       brd_content: str = None,
                       data_models: Dict = None) -> io.BytesIO:
    """
    Create an Excel file with multiple sheets.
    
    Args:
        dataframes: Dictionary of {sheet_name: DataFrame}
        include_brd: Whether to include BRD sheet
        brd_content: BRD content to include
        data_models: Dictionary of data model info
        
    Returns:
        BytesIO object containing the Excel file
    """
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Add BRD sheet first if available
        if include_brd and brd_content:
            brd_df = pd.DataFrame({'Business Requirements Document': [brd_content]})
            brd_df.to_excel(writer, sheet_name='BRD', index=False)
            
            # Format BRD sheet
            try:
                from openpyxl.styles import Alignment, Font
                brd_ws = writer.sheets['BRD']
                brd_ws.column_dimensions['A'].width = 120
                
                for row in brd_ws.iter_rows():
                    for cell in row:
                        cell.alignment = Alignment(wrap_text=True, vertical='top')
                
                brd_ws['A1'].font = Font(bold=True, size=12)
                
                content_lines = brd_content.count('\n') + 1
                brd_ws.row_dimensions[2].height = min(content_lines * 15, 600)
            except Exception:
                pass
        
        # Add data sheets
        for sheet_name, df in dataframes.items():
            if df is not None and len(df) > 0:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        # Add data model sheets if available
        if data_models:
            _add_data_model_sheets(writer, data_models)
    
    output.seek(0)
    return output


def _add_data_model_sheets(writer, data_models: Dict):
    """Add data model images to Excel workbook."""
    try:
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Font
        import tempfile
        import os
        
        workbook = writer.book
        temp_files = []
        
        for model_type in ['conceptual', 'logical', 'physical']:
            if model_type not in data_models:
                continue
                
            model_data = data_models[model_type]
            
            if model_data.get('png'):
                sheet_name = f'DM_{model_type[:4].title()}'
                ws = workbook.create_sheet(title=sheet_name)
                
                ws['A1'] = f'{model_type.title()} Data Model'
                ws['A1'].font = Font(bold=True, size=14)
                ws['A3'] = 'Source-to-target data flow diagram'
                
                # Save PNG to temp file
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                tmp.write(model_data['png'])
                tmp.close()
                temp_files.append(tmp.name)
                
                try:
                    img = XLImage(tmp.name)
                    # Scale image
                    if img.width > 800:
                        ratio = 800 / img.width
                        img.width = 800
                        img.height = int(img.height * ratio)
                    if img.height > 600:
                        ratio = 600 / img.height
                        img.height = 600
                        img.width = int(img.width * ratio)
                    ws.add_image(img, 'A5')
                except Exception:
                    ws['A5'] = 'Image could not be embedded'
            
            elif model_data.get('dot_source'):
                sheet_name = f'DM_{model_type[:4].title()}'
                ws = workbook.create_sheet(title=sheet_name)
                ws['A1'] = f'{model_type.title()} Data Model (DOT Source)'
                ws['A1'].font = Font(bold=True, size=14)
                ws['A3'] = 'Graphviz DOT source - paste into https://graphviz.org'
                ws['A5'] = model_data['dot_source']
        
        # Cleanup temp files
        for tmp_path in temp_files:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
                
    except ImportError:
        pass
    except Exception:
        pass


def create_word_export(content: str,
                      title_text: str = "Business Requirements Document (BRD)",
                      mapping_name: str = "mapping",
                      model_info: str = "") -> Optional[io.BytesIO]:
    """
    Create a Word document from markdown content.

    Args:
        content: Markdown content to convert
        title_text: Document title (e.g. "Technical Design Document (TDD)")
        mapping_name: Name of the mapping for metadata
        model_info: Model information string

    Returns:
        BytesIO object containing the Word file, or None if docx not available
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()

    # Add title
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    import datetime
    doc.add_paragraph(f"Mapping: {mapping_name}")
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if model_info:
        doc.add_paragraph(f"Model: {model_info}")
    doc.add_paragraph("")

    # Parse and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            doc.add_paragraph(line)

    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer


def create_sql_export(sql_content: str, filename: str = "generated_sql.sql") -> io.BytesIO:
    """
    Create a downloadable SQL file.
    
    Args:
        sql_content: SQL content
        filename: Suggested filename
        
    Returns:
        BytesIO object containing the SQL file
    """
    output = io.BytesIO()
    output.write(sql_content.encode('utf-8'))
    output.seek(0)
    return output
